"""Template merge — <<Placeholder>> substitution via python-docx."""
import re
from pathlib import Path

from docx import Document


def build_replacements(row, config, template_fields=None):
    """Build the placeholder replacement mapping for one data row.

    Matches template placeholders against data columns and config fields.
    Column matching is case-insensitive with spaces/underscores normalized.

    Args:
        row: dict from one DataFrame row (e.g. {"Learner Name": "Tan Wei Ming", "Grades": "Distinction"})
        config: {"programme_name": str, "start_date": str, "end_date": str, ...}
        template_fields: list of placeholder names from the template (e.g. ["Learner Name", "Grades"]).
            If None, falls back to matching common fields.

    Returns:
        Dict mapping placeholder names to replacement values.
        Keys match <<Placeholder>> text in the template.
    """
    replacements = {}

    start = config.get("start_date", "")
    end = config.get("end_date", "")
    if start and end:
        programme_date = f"{start} to {end}"
    elif end:
        programme_date = end
    else:
        programme_date = start

    # Config-derived fields (entered by the user, not from the spreadsheet).
    # Each value lists every template placeholder spelling that should resolve
    # to it, so a template can say <<Programme Start Date>> or <<Start Date>>.
    config_fields = {
        "Programme Name": config.get("programme_name", ""),
        "Programme Start Date": start,
        "Start Date": start,
        "Programme End Date": end,
        "End Date": end,
        "Programme Date": programme_date,
    }

    if template_fields:
        # Dynamic matching: for each template placeholder, try the user-entered
        # config fields FIRST, then fall back to spreadsheet columns. Programme
        # name/dates are programme-wide constants the user typed in, so they must
        # take precedence over any same-named column in the spreadsheet.
        for field in template_fields:
            field_norm = field.lower().replace(" ", "").replace("_", "")

            # Try config fields first (user input is authoritative)
            matched = False
            for cf_name, cf_val in config_fields.items():
                cf_norm = cf_name.lower().replace(" ", "").replace("_", "")
                if cf_norm == field_norm:
                    replacements[field] = cf_val
                    matched = True
                    break

            # If not a config field, try spreadsheet columns
            if not matched:
                for col, val in row.items():
                    col_norm = col.lower().replace(" ", "").replace("_", "")
                    if col_norm == field_norm:
                        if val is None or (isinstance(val, float) and str(val) == "nan"):
                            replacements[field] = ""
                        else:
                            replacements[field] = str(val)
                        matched = True
                        break

            # If still not matched, use empty string
            if not matched:
                replacements[field] = ""
    else:
        # Fallback: match all data columns + config fields
        for col, val in row.items():
            if val is None or (isinstance(val, float) and str(val) == "nan"):
                replacements[col] = ""
            else:
                replacements[col] = str(val)
        replacements.update(config_fields)

    return replacements


def merge_template(template_path, replacements, output_path):
    """Merge a .docx template by replacing <<Placeholder>> text.

    Args:
        template_path: Path to the .docx template
        replacements: Dict mapping placeholder names to values
        output_path: Where to save the merged document

    Returns:
        The output_path string

    Raises:
        FileNotFoundError: if template does not exist
        ValueError: if template has no placeholders
    """
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template))

    placeholders = _extract_placeholders(doc)
    if not placeholders:
        raise ValueError(
            f"Template '{template_path}' has no placeholders. "
            "Add <<Learner Name>>, <<Grades>>, etc. in your template."
        )

    applicable = {k: v for k, v in replacements.items() if k in placeholders}
    _replace_placeholders(doc, applicable)

    doc.save(output_path)
    return output_path


def _extract_placeholders(doc):
    """Find all <<Placeholder>> patterns in a python-docx Document."""
    placeholders = set()
    for paragraph in doc.paragraphs:
        placeholders.update(m.strip() for m in re.findall(r'<<(.+?)>>', paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    placeholders.update(m.strip() for m in re.findall(r'<<(.+?)>>', paragraph.text))
    for section in doc.sections:
        for part in [section.header, section.first_page_header, section.even_page_header,
                      section.footer, section.first_page_footer, section.even_page_footer]:
            if part and not getattr(part, 'is_linked_to_previous', True):
                for paragraph in part.paragraphs:
                    placeholders.update(m.strip() for m in re.findall(r'<<(.+?)>>', paragraph.text))
    return placeholders


def _replace_placeholders(doc, replacements):
    """Replace <<Placeholder>> with values throughout the document."""
    replace_map = {}
    for key, value in replacements.items():
        replace_map[f'<<{key}>>'] = str(value)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replace_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replace_map)

    for section in doc.sections:
        for part in [section.header, section.first_page_header, section.even_page_header,
                      section.footer, section.first_page_footer, section.even_page_footer]:
            if part and not getattr(part, 'is_linked_to_previous', True):
                for paragraph in part.paragraphs:
                    _replace_in_paragraph(paragraph, replace_map)


def _replace_cross_run(runs, pattern, value):
    """Replace a placeholder that spans multiple runs, preserving formatting.

    Strategy: concatenate all run texts, find the pattern, then distribute
    the replacement back across the original runs — keeping each run's
    formatting (bold, italic, font, size, etc.) intact.
    """
    texts = [r.text for r in runs]
    combined = ''.join(texts)

    if pattern not in combined:
        return

    # Build a character-offset map: char index -> (run index, offset within run)
    run_boundaries = []
    for i, t in enumerate(texts):
        run_boundaries.append((i, len(t)))

    # Find pattern position in combined text
    start = combined.index(pattern)
    end = start + len(pattern)

    # Map combined-text offsets to run indices and positions within those runs
    char_count = 0
    start_run = start_pos = end_run = end_pos = None
    for run_idx, run_len in run_boundaries:
        run_start = char_count
        run_end = char_count + run_len

        if start_run is None and start < run_end:
            start_run = run_idx
            start_pos = start - run_start

        if end_run is None and end <= run_end:
            end_run = run_idx
            end_pos = end - run_start

        char_count = run_end

    if start_run is None or end_run is None:
        return

    if start_run == end_run:
        # Pattern is within a single run (shouldn't reach here, but safe fallback)
        runs[start_run].text = (
            runs[start_run].text[:start_pos] + value + runs[start_run].text[end_pos:]
        )
    else:
        # Pattern spans multiple runs
        # First run: keep text before the pattern start, prepend the value
        runs[start_run].text = runs[start_run].text[:start_pos] + value

        # Intermediate runs (between first and last): clear entirely
        for i in range(start_run + 1, end_run):
            runs[i].text = ''

        # Last run: keep text after the pattern end
        runs[end_run].text = runs[end_run].text[end_pos:]


def _replace_in_paragraph(paragraph, replace_map):
    """Replace placeholders in a paragraph, handling cross-run cases."""
    full_text = paragraph.text
    if '<<' not in full_text:
        return

    # Try single-run replacement first (preserves formatting)
    for run in paragraph.runs:
        for pattern, value in replace_map.items():
            if pattern in run.text:
                run.text = run.text.replace(pattern, value)

    # If all placeholders replaced, done
    if '<<' not in paragraph.text:
        return

    # Cross-run case: reconstruct runs preserving each run's formatting
    runs = paragraph.runs
    if not runs:
        return

    for pattern, value in replace_map.items():
        _replace_cross_run(runs, pattern, value)


def get_template_fields(template_path):
    """Return the set of placeholder names detected in a template.

    Raises:
        FileNotFoundError: if template does not exist
    """
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template))
    return _extract_placeholders(doc)
