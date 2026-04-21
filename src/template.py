"""Template validation and placeholder detection via python-docx."""
import re
import tempfile
from pathlib import Path

from docx import Document


def _extract_placeholders(doc):
    """Find all <<Placeholder>> patterns in a python-docx Document.

    Scans paragraphs (including tables, headers, footers).
    """
    placeholders = set()
    for paragraph in doc.paragraphs:
        placeholders.update(m.strip() for m in re.findall(r'<<(.+?)>>', paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    placeholders.update(m.strip() for m in re.findall(r'<<(.+?)>>', paragraph.text))
    for section in doc.sections:
        for part in [section.header, section.first_page_header, section.even_page_header,
                      section.footer, section.first_page_footer, section.even_page_footer]:
            if part and not getattr(part, 'is_linked_to_previous', True):
                for paragraph in part.paragraphs:
                    placeholders.update(m.strip() for m in re.findall(r'<<(.+?)>>', paragraph.text))
    return placeholders


def detect_merge_fields(template_path):
    """Detect <<Placeholder>> entries in a Word template.

    Returns sorted list of unique placeholder names.
    """
    doc = Document(template_path)
    fields = _extract_placeholders(doc)
    return sorted(fields)


def validate_template(uploaded_file):
    """Validate an uploaded Word template.

    Returns:
        {"valid": bool, "fields": list[str], "error": str|None}
    """
    if uploaded_file is None:
        return {"valid": False, "fields": [], "error": "No file uploaded."}

    name = getattr(uploaded_file, 'name', '') or ''
    if not name.lower().endswith('.docx'):
        return {"valid": False, "fields": [], "error": "Please upload a Word document (.docx)."}

    tmp = None
    try:
        content = uploaded_file.getvalue()
        if len(content) == 0:
            return {"valid": False, "fields": [], "error": "The file is empty. Please upload a valid .docx file."}

        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.write(content)
        tmp.close()

        doc = Document(tmp.name)
        fields = _extract_placeholders(doc)

        if not fields:
            return {
                "valid": False,
                "fields": [],
                "error": "No placeholders found. Add <<Learner Name>>, <<Grades>>, <<Programme Name>>, <<End Date>> in your template."
            }

        return {"valid": True, "fields": sorted(fields), "error": None}

    except Exception:
        return {"valid": False, "fields": [], "error": "Could not read this file. It may be corrupted or not a valid Word document."}

    finally:
        if tmp:
            try:
                Path(tmp.name).unlink(missing_ok=True)
            except Exception:
                pass
