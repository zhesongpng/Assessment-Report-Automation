"""Shared test fixtures."""
import os
import pytest
import tempfile
from pathlib import Path
from docx import Document

SAMPLE_DATA = os.path.join(
    os.path.dirname(__file__), "..",
    "workspaces", "assessment-automation", "briefs",
    "Assessment Result - For Release.xlsx"
)


@pytest.fixture
def sample_template_path(tmp_path):
    """Create a test template with <<...>> placeholders."""
    doc = Document()
    doc.add_paragraph("Assessment Report")
    doc.add_paragraph("Learner: <<Learner Name>>")
    doc.add_paragraph("Grade: <<Grades>>")
    doc.add_paragraph("Programme: <<Programme Name>>")
    doc.add_paragraph("Date: <<End Date>>")
    template_path = str(tmp_path / "test_template.docx")
    doc.save(template_path)
    return template_path


@pytest.fixture
def sample_data_path():
    return SAMPLE_DATA
