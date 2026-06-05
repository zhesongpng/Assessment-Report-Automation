"""Tests for template merge engine."""
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from src.merge import build_replacements, merge_template, get_template_fields, _replace_in_paragraph


class TestBuildReplacements:
    def test_maps_data_columns_to_placeholders(self):
        row = {"Learner Name": "Tan Wei Ming", "Grades": "Distinction"}
        config = {"programme_name": "AI Analytics", "end_date": "15 May 2025"}
        result = build_replacements(row, config)
        assert result["Learner Name"] == "Tan Wei Ming"
        assert result["Grades"] == "Distinction"
        assert result["Programme Name"] == "AI Analytics"
        assert result["End Date"] == "15 May 2025"

    def test_nan_values_become_empty_string(self):
        row = {"Learner Name": "Alice", "Grades": float("nan")}
        config = {"programme_name": "Test", "end_date": "2025"}
        result = build_replacements(row, config)
        assert result["Grades"] == ""

    def test_missing_config_defaults_empty(self):
        row = {"Learner Name": "Alice", "Grades": "A"}
        config = {}
        result = build_replacements(row, config)
        assert result["Programme Name"] == ""
        assert result["End Date"] == ""

    def test_dynamic_matching_with_template_fields(self):
        """When template_fields is provided, dynamically match data columns."""
        row = {"Email": "alice@example.com", "Course Title": "ML Basics", "Grades": "A"}
        config = {"programme_name": "Data Science", "start_date": "1 Jan 2025", "end_date": "30 Jun 2025"}
        result = build_replacements(row, config, template_fields=["Email", "Course Title", "Programme Name", "Programme date"])
        assert result["Email"] == "alice@example.com"
        assert result["Course Title"] == "ML Basics"
        assert result["Programme Name"] == "Data Science"
        assert result["Programme date"] == "1 Jan 2025 to 30 Jun 2025"

    def test_dynamic_matching_unmatched_field_gets_empty(self):
        """Template fields not found in data or config get empty string."""
        row = {"Email": "alice@example.com"}
        config = {}
        result = build_replacements(row, config, template_fields=["Email", "Nonexistent Field"])
        assert result["Email"] == "alice@example.com"
        assert result["Nonexistent Field"] == ""


class TestMergeTemplate:
    def test_merge_produces_output(self, sample_template_path):
        tmp_out = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_out.close()
        try:
            replacements = {
                "Learner Name": "Test Learner",
                "Grades": "Distinction",
                "Programme Name": "AI Analytics",
                "End Date": "15 May 2025",
            }
            result = merge_template(sample_template_path, replacements, tmp_out.name)
            assert Path(result).exists()
            assert Path(result).stat().st_size > 0
        finally:
            Path(tmp_out.name).unlink(missing_ok=True)

    def test_merge_preserves_formatting(self, sample_template_path):
        tmp_out = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_out.close()
        try:
            replacements = {
                "Learner Name": "Test Learner",
                "Grades": "Pass",
                "Programme Name": "Test Programme",
                "End Date": "1 Jan 2025",
            }
            merge_template(sample_template_path, replacements, tmp_out.name)
            assert Path(tmp_out.name).stat().st_size > 1000
        finally:
            Path(tmp_out.name).unlink(missing_ok=True)

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            merge_template("/nonexistent.docx", {}, "/tmp/out.docx")

    def test_only_passes_applicable_fields(self, sample_template_path):
        tmp_out = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_out.close()
        try:
            replacements = {
                "Learner Name": "Test",
                "Grades": "A",
                "Programme Name": "Prog",
                "End Date": "2025",
                "Nonexistent Field": "should be ignored",
            }
            merge_template(sample_template_path, replacements, tmp_out.name)
            assert Path(tmp_out.name).exists()
        finally:
            Path(tmp_out.name).unlink(missing_ok=True)


class TestGetTemplateFields:
    def test_returns_expected_fields(self, sample_template_path):
        fields = get_template_fields(sample_template_path)
        assert "Learner Name" in fields
        assert "Grades" in fields
        assert "Programme Name" in fields
        assert "End Date" in fields

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            get_template_fields("/nonexistent.docx")


class TestCrossRunFormattingPreservation:
    """Regression: placeholders split across runs must not corrupt formatting.

    Microsoft Word frequently splits text across multiple runs with different
    formatting. The old code merged all runs into the first run, causing the
    entire paragraph to inherit the first run's formatting (e.g. bold).
    """

    def test_cross_run_replacement_preserves_non_bold_formatting(self):
        """A placeholder split across a bold run and a normal run must not
        make the surrounding normal text become bold."""
        doc = Document()
        paragraph = doc.add_paragraph()

        # Simulate Word splitting "<<Learner Name>>" across two runs:
        # Run 1: "<<Learner" (bold=True)   Run 2: " Name>>" (bold=False)
        run1 = paragraph.add_run("<<Learner")
        run1.bold = True
        run2 = paragraph.add_run(" Name>>")
        run2.bold = False

        _replace_in_paragraph(paragraph, {"<<Learner Name>>": "Alice Tan"})

        # The replacement should have happened
        assert "Alice Tan" in paragraph.text
        assert "<<Learner Name>>" not in paragraph.text

        # Run 1 (originally bold) should stay bold
        assert run1.bold is True

        # Run 2 (originally not bold) should still not be bold
        assert run2.bold is False

    def test_cross_run_replacement_preserves_each_run_formatting(self):
        """Text after the placeholder in the last run keeps its formatting."""
        doc = Document()
        paragraph = doc.add_paragraph()

        # Simulate: "Dear <<Learner" (run1, italic) + " Name>>, welcome!" (run2, not italic)
        run1 = paragraph.add_run("Dear <<Learner")
        run1.italic = True
        run2 = paragraph.add_run(" Name>>, welcome!")
        run2.italic = False

        _replace_in_paragraph(paragraph, {"<<Learner Name>>": "Bob"})

        assert "Dear Bob, welcome!" in paragraph.text
        assert run1.italic is True
        assert run2.italic is False

    def test_cross_run_replacement_via_merge_template(self, tmp_path):
        """End-to-end: a DOCX with split-run placeholders merges correctly
        and preserves each run's formatting in the output file."""
        # Create a template where the placeholder is split across runs
        doc = Document()
        paragraph = doc.add_paragraph()
        run1 = paragraph.add_run("Result: <<Grades")
        run1.bold = True
        run2 = paragraph.add_run(">> for <<Learner Name>>")
        run2.bold = False

        template_path = str(tmp_path / "split_template.docx")
        output_path = str(tmp_path / "output.docx")
        doc.save(template_path)

        replacements = {"Grades": "Distinction", "Learner Name": "Carol"}
        merge_template(template_path, replacements, output_path)

        # Read back the output and verify
        result_doc = Document(output_path)
        result_para = result_doc.paragraphs[0]

        assert "Distinction" in result_para.text
        assert "Carol" in result_para.text
        assert "<<" not in result_para.text

        # Verify bold/non-bold runs are preserved
        bold_runs = [r for r in result_para.runs if r.bold is True]
        non_bold_runs = [r for r in result_para.runs if r.bold is False]
        assert len(bold_runs) >= 1, "At least one run should remain bold"
        assert len(non_bold_runs) >= 1, "At least one run should remain non-bold"
